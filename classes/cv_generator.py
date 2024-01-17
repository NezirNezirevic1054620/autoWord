from docx import Document


class CVGenerator:
    def __init__(self, title, paragraph):
        self.title = title
        self.paragraph = paragraph

    @staticmethod
    def __get_document():
        document = Document()

        return document

    def generate_cv(self):
        document = self.__get_document()
        document.add_heading(self.title, 0)
        document.add_paragraph(self.paragraph)
        document.save(f"docs/{self.title}.docx")